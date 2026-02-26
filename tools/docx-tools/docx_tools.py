import sys
import argparse
import json
try:
    import docx
except ImportError:
    print("python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

def extract_content(file_path):
    """Extracts text and tables from a DOCX file."""
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening DOCX: {e}")
        sys.exit(1)
        
    content = {
        "paragraphs": [p.text for p in doc.paragraphs if p.text.strip()],
        "tables": []
    }
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content["tables"].append(table_data)
        
    print(json.dumps(content, indent=2))

def create_document(output_path, title, paragraphs):
    """Creates a simple new DOCX file."""
    doc = docx.Document()
    doc.add_heading(title, 0)
    
    for p in paragraphs:
        doc.add_paragraph(p)
        
    doc.save(output_path)
    print(f"Created document at {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="DOCX Manipulation Tool")
    subparsers = parser.add_subparsers(dest="command", required=True)
    
    # Extract
    ext_parser = subparsers.add_parser("extract", help="Extract text and tables as JSON")
    ext_parser.add_argument("file", help="Path to DOCX file")
    
    # Create
    cre_parser = subparsers.add_parser("create", help="Create a straightforward DOCX")
    cre_parser.add_argument("output", help="Output path for the DOCX")
    cre_parser.add_argument("--title", required=True, help="Document title")
    cre_parser.add_argument("--text", nargs="+", required=True, help="Paragraphs of text")
    
    args = parser.parse_args()
    
    if args.command == "extract":
        extract_content(args.file)
    elif args.command == "create":
        create_document(args.output, args.title, args.text)
